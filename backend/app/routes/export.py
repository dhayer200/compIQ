"""Multi-format report export: Markdown, DOCX, PDF."""
import io
import json
from uuid import UUID

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from app.db import get_pool
from app.models.report import RentEstimate

router = APIRouter()


def _fmt(n: int) -> str:
    return f"${n:,}"


async def _load_analysis(analysis_id: UUID) -> dict:
    """Load full analysis data from DB."""
    db = await get_pool()

    row = await db.fetchrow("SELECT * FROM analyses WHERE id = $1", analysis_id)
    if not row:
        raise HTTPException(status_code=404, detail="Analysis not found")

    prop = await db.fetchrow(
        "SELECT * FROM properties WHERE id = $1", row["subject_property_id"]
    )
    comp_rows = await db.fetch(
        "SELECT * FROM comp_properties WHERE analysis_id = $1 ORDER BY correlation DESC",
        analysis_id,
    )

    return {"row": row, "prop": prop, "comp_rows": comp_rows}


def _build_markdown(data: dict) -> str:
    """Generate a full Markdown report."""
    row = data["row"]
    prop = data["prop"]
    comps = data["comp_rows"]

    lines = []
    lines.append(f"# Comparative Market Analysis")
    lines.append(f"**{prop['address']}**")
    if prop["city"] or prop["state"]:
        lines.append(f"{prop['city'] or ''}, {prop['state'] or ''} {prop['zip_code'] or ''}")
    lines.append("")

    # Subject property
    lines.append("## Subject Property")
    lines.append("")
    details = []
    if prop["bedrooms"] is not None:
        details.append(f"- **Bedrooms:** {prop['bedrooms']}")
    if prop["bathrooms"] is not None:
        details.append(f"- **Bathrooms:** {float(prop['bathrooms'])}")
    if prop["sqft"]:
        details.append(f"- **Sq Ft:** {prop['sqft']:,}")
    if prop["lot_size"]:
        details.append(f"- **Lot Size:** {prop['lot_size']:,} sq ft")
    if prop["year_built"]:
        details.append(f"- **Year Built:** {prop['year_built']}")
    if prop["property_type"]:
        details.append(f"- **Type:** {prop['property_type']}")
    if prop["last_sale_date"]:
        details.append(f"- **Last Sale Date:** {prop['last_sale_date']}")
    if prop["last_sale_price"]:
        details.append(f"- **Last Sale Price:** {_fmt(prop['last_sale_price'])}")
    lines.extend(details)
    lines.append("")

    # Valuation
    lines.append("## Valuation Summary")
    lines.append("")
    lines.append(f"| Metric | Value |")
    lines.append(f"|--------|-------|")
    lines.append(f"| **Estimated Value** | **{_fmt(row['estimated_value'])}** |")
    lines.append(f"| Range Low | {_fmt(row['value_range_low'])} |")
    lines.append(f"| Range High | {_fmt(row['value_range_high'])} |")
    lines.append(f"| Monthly Rent | {_fmt(row['estimated_rent'])} |")
    lines.append(f"| Rent Range | {_fmt(row['rent_range_low'])} - {_fmt(row['rent_range_high'])} |")
    if row["est_days_on_market"]:
        lines.append(f"| Avg Days on Market | {row['est_days_on_market']} |")
    lines.append("")

    # Cash flow potential
    if comps:
        prices = [r["sale_price"] for r in comps if r["sale_price"]]
        if prices:
            lines.append("## Current Cash Flow Potential")
            lines.append("")
            lines.append(f"- **Lowest Comp:** {_fmt(min(prices))}")
            lines.append(f"- **Highest Comp:** {_fmt(max(prices))}")
            lines.append(f"- **Comp Spread:** {_fmt(max(prices) - min(prices))}")
            lines.append("")

    # Comps table
    lines.append(f"## Comparable Sales ({len(comps)})")
    lines.append("")
    lines.append("| Address | Sale Price | Date | Bed/Bath | Sq Ft | Distance | Adjusted |")
    lines.append("|---------|-----------|------|----------|-------|----------|----------|")
    for c in comps:
        beds = c["bedrooms"] or "-"
        baths = float(c["bathrooms"]) if c["bathrooms"] else "-"
        sqft = f"{c['sqft']:,}" if c["sqft"] else "-"
        dist = f"{float(c['distance_miles']):.2f}mi" if c["distance_miles"] else "-"
        adj = _fmt(c["adjusted_price"]) if c["adjusted_price"] else "-"
        dt = str(c["sale_date"]) if c["sale_date"] else "-"
        lines.append(f"| {c['address']} | {_fmt(c['sale_price'])} | {dt} | {beds}/{baths} | {sqft} | {dist} | {adj} |")
    lines.append("")

    # Adjustments
    has_adj = any(c["adjustments"] for c in comps)
    if has_adj:
        lines.append("## Adjustments Detail")
        lines.append("")
        for c in comps:
            if c["adjustments"]:
                adj = json.loads(c["adjustments"]) if isinstance(c["adjustments"], str) else c["adjustments"]
                nonzero = {k: v for k, v in adj.items() if v != 0}
                if nonzero:
                    lines.append(f"**{c['address']}** (Sale: {_fmt(c['sale_price'])})")
                    for k, v in nonzero.items():
                        lines.append(f"- {k}: {_fmt(v) if v >= 0 else '-' + _fmt(abs(v))}")
                    lines.append(f"- **Adjusted Price: {_fmt(c['adjusted_price'])}**")
                    lines.append("")

    # Narrative
    if row["narrative_text"]:
        lines.append("## Analysis Narrative")
        lines.append("")
        lines.append(row["narrative_text"])
        lines.append("")

    # Disclaimer
    lines.append("---")
    lines.append("*This report is generated by compIQ and is intended for informational purposes only. "
                 "It is not a formal appraisal and should not be used as a substitute for a licensed appraisal.*")

    return "\n".join(lines)


def _build_docx(data: dict) -> io.BytesIO:
    """Generate a DOCX report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    row = data["row"]
    prop = data["prop"]
    comps = data["comp_rows"]

    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Comparative Market Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subject address
    addr_para = doc.add_paragraph()
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = addr_para.add_run(prop["address"])
    run.font.size = Pt(14)
    run.bold = True
    if prop["city"] or prop["state"]:
        addr_para.add_run(f"\n{prop['city'] or ''}, {prop['state'] or ''} {prop['zip_code'] or ''}")

    doc.add_paragraph("")

    # Subject property details
    doc.add_heading("Subject Property", level=1)
    details = []
    if prop["bedrooms"] is not None:
        details.append(("Bedrooms", str(prop["bedrooms"])))
    if prop["bathrooms"] is not None:
        details.append(("Bathrooms", str(float(prop["bathrooms"]))))
    if prop["sqft"]:
        details.append(("Square Feet", f"{prop['sqft']:,}"))
    if prop["lot_size"]:
        details.append(("Lot Size", f"{prop['lot_size']:,} sq ft"))
    if prop["year_built"]:
        details.append(("Year Built", str(prop["year_built"])))
    if prop["property_type"]:
        details.append(("Property Type", prop["property_type"]))
    if prop["last_sale_date"]:
        details.append(("Last Sale Date", str(prop["last_sale_date"])))
    if prop["last_sale_price"]:
        details.append(("Last Sale Price", _fmt(prop["last_sale_price"])))

    if details:
        table = doc.add_table(rows=len(details), cols=2)
        table.style = "Light Grid Accent 1"
        for i, (label, value) in enumerate(details):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

    doc.add_paragraph("")

    # Valuation
    doc.add_heading("Valuation Summary", level=1)
    val_data = [
        ("Estimated Value", _fmt(row["estimated_value"])),
        ("Value Range", f"{_fmt(row['value_range_low'])} - {_fmt(row['value_range_high'])}"),
        ("Monthly Rent", _fmt(row["estimated_rent"])),
        ("Rent Range", f"{_fmt(row['rent_range_low'])} - {_fmt(row['rent_range_high'])}"),
    ]
    if row["est_days_on_market"]:
        val_data.append(("Avg Days on Market", str(row["est_days_on_market"])))

    table = doc.add_table(rows=len(val_data), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(val_data):
        table.rows[i].cells[0].text = label
        cell = table.rows[i].cells[1]
        cell.text = value
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    doc.add_paragraph("")

    # Cash flow potential
    if comps:
        prices = [r["sale_price"] for r in comps if r["sale_price"]]
        if prices:
            doc.add_heading("Current Cash Flow Potential", level=1)
            doc.add_paragraph(f"Comp Price Range: {_fmt(min(prices))} - {_fmt(max(prices))}")
            doc.add_paragraph(f"Spread: {_fmt(max(prices) - min(prices))}")
            doc.add_paragraph("")

    # Comps table
    doc.add_heading(f"Comparable Sales ({len(comps)})", level=1)
    headers = ["Address", "Sale Price", "Date", "Bed/Bath", "Sq Ft", "Dist", "Adjusted"]
    table = doc.add_table(rows=1 + len(comps), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, c in enumerate(comps, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = c["address"] or ""
        row_cells[1].text = _fmt(c["sale_price"])
        row_cells[2].text = str(c["sale_date"]) if c["sale_date"] else "-"
        beds = c["bedrooms"] or "-"
        baths = float(c["bathrooms"]) if c["bathrooms"] else "-"
        row_cells[3].text = f"{beds}/{baths}"
        row_cells[4].text = f"{c['sqft']:,}" if c["sqft"] else "-"
        row_cells[5].text = f"{float(c['distance_miles']):.2f}mi" if c["distance_miles"] else "-"
        row_cells[6].text = _fmt(c["adjusted_price"]) if c["adjusted_price"] else "-"

    doc.add_paragraph("")

    # Narrative
    if row["narrative_text"]:
        doc.add_heading("Analysis Narrative", level=1)
        for para in row["narrative_text"].split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Disclaimer
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "This report is generated by compIQ and is intended for informational purposes only. "
        "It is not a formal appraisal and should not be used as a substitute for a licensed appraisal."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/analysis/{analysis_id}/export/md")
async def export_markdown(analysis_id: UUID):
    data = await _load_analysis(analysis_id)
    md = _build_markdown(data)
    address = data["prop"]["address"].replace(",", "").replace(" ", "_")[:40]
    return StreamingResponse(
        io.BytesIO(md.encode("utf-8")),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="CMA_{address}.md"'},
    )


@router.get("/analysis/{analysis_id}/export/docx")
async def export_docx(analysis_id: UUID):
    data = await _load_analysis(analysis_id)
    buf = _build_docx(data)
    address = data["prop"]["address"].replace(",", "").replace(" ", "_")[:40]
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="CMA_{address}.docx"'},
    )
